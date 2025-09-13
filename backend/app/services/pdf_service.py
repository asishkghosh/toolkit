import os
import tempfile
import uuid
from typing import List, Optional
import logging
from PyPDF2 import PdfReader, PdfWriter
from pdf2docx import Converter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import io
import fitz  # PyMuPDF for advanced compression

# Handle pdfplumber import with fallback
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available. Install it for better page counting.")

logger = logging.getLogger(__name__)

class PDFService:
    """Service class for PDF operations"""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    async def merge_pdfs(self, pdf_paths: List[str]) -> str:
        """Merge multiple PDF files into one"""
        try:
            writer = PdfWriter()
            
            # Add all pages from all PDFs
            for pdf_path in pdf_paths:
                with open(pdf_path, 'rb') as file:
                    reader = PdfReader(file)
                    for page in reader.pages:
                        writer.add_page(page)
            
            # Generate output path
            output_path = os.path.join(self.temp_dir, f"merged_{uuid.uuid4().hex}.pdf")
            
            # Write merged PDF
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            logger.info(f"Successfully merged {len(pdf_paths)} PDFs to {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error merging PDFs: {str(e)}")
            raise
    
    async def get_page_count(self, pdf_path: str) -> int:
        """Get the accurate page count of a PDF file using multiple methods for reliability"""
        page_count = None
        methods_tried = []
        
        # Method 1: PyPDF2 PdfReader
        try:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                page_count = len(reader.pages)
                methods_tried.append(f"PyPDF2: {page_count} pages")
                logger.info(f"PyPDF2 page count: {page_count} pages")
                
                # Validate the result makes sense
                if page_count > 0 and page_count < 10000:  # Reasonable bounds
                    return page_count
                    
        except Exception as e:
            methods_tried.append(f"PyPDF2 failed: {str(e)}")
            logger.warning(f"PyPDF2 method failed: {str(e)}")
        
        # Method 2: Try with pdfplumber (already in requirements) if available
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    page_count = len(pdf.pages)
                    methods_tried.append(f"pdfplumber: {page_count} pages")
                    logger.info(f"pdfplumber page count: {page_count} pages")
                    
                    if page_count > 0 and page_count < 10000:
                        return page_count
                        
            except Exception as e:
                methods_tried.append(f"pdfplumber failed: {str(e)}")
                logger.warning(f"pdfplumber method failed: {str(e)}")
        else:
            methods_tried.append("pdfplumber not available")
        
        # Method 3: Try with fitz (PyMuPDF) if available
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(pdf_path)
            page_count = doc.page_count
            doc.close()
            methods_tried.append(f"PyMuPDF: {page_count} pages")
            logger.info(f"PyMuPDF page count: {page_count} pages")
            
            if page_count > 0 and page_count < 10000:
                return page_count
                
        except ImportError:
            methods_tried.append("PyMuPDF not available")
            pass
        except Exception as e:
            methods_tried.append(f"PyMuPDF failed: {str(e)}")
            logger.warning(f"PyMuPDF method failed: {str(e)}")
        
        # Method 4: Parse PDF manually for page count
        try:
            with open(pdf_path, 'rb') as file:
                content = file.read()
                
                # Try to find the /Count entry in the /Pages dictionary
                import re
                pages_pattern = r'/Type\s*/Pages[^}]*?/Count\s+(\d+)'
                match = re.search(pages_pattern, content.decode('latin-1', errors='ignore'), re.IGNORECASE | re.DOTALL)
                
                if match:
                    page_count = int(match.group(1))
                    methods_tried.append(f"Manual parsing: {page_count} pages")
                    logger.info(f"Manual parsing page count: {page_count} pages")
                    
                    if page_count > 0 and page_count < 10000:
                        return page_count
                
                # Fallback: count individual page objects
                page_objects = re.findall(r'/Type\s*/Page(?![s\w])', content.decode('latin-1', errors='ignore'), re.IGNORECASE)
                if page_objects:
                    page_count = len(page_objects)
                    methods_tried.append(f"Page object count: {page_count} pages")
                    logger.info(f"Page object counting: {page_count} pages")
                    
                    if page_count > 0 and page_count < 10000:
                        return page_count
                        
        except Exception as e:
            methods_tried.append(f"Manual parsing failed: {str(e)}")
            logger.warning(f"Manual parsing failed: {str(e)}")
        
        # Method 5: Try with different PyPDF2 approach
        try:
            from PyPDF2 import PdfFileReader
            with open(pdf_path, 'rb') as file:
                reader = PdfFileReader(file, strict=False)
                page_count = reader.getNumPages()
                methods_tried.append(f"PyPDF2 legacy: {page_count} pages")
                logger.info(f"PyPDF2 legacy page count: {page_count} pages")
                
                if page_count > 0 and page_count < 10000:
                    return page_count
                    
        except Exception as e:
            methods_tried.append(f"PyPDF2 legacy failed: {str(e)}")
            logger.warning(f"PyPDF2 legacy method failed: {str(e)}")
        
        # If all methods failed or returned unreasonable results
        logger.error(f"All page counting methods failed or returned invalid results. Methods tried: {methods_tried}")
        
        # Return a minimal fallback
        if page_count and page_count > 0:
            logger.warning(f"Returning last valid count: {page_count}")
            return page_count
        
        # Final fallback
        logger.warning("Returning fallback page count of 1")
        return 1
    
    async def split_at_page(self, pdf_path: str, split_page: int) -> List[str]:
        """Split PDF into two parts at the specified page number with preserved formatting"""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                total_pages = len(reader.pages)
                
                # Validate split page number
                if split_page is None or split_page < 1:
                    raise ValueError("Split page number must be a positive integer")
                
                if split_page >= total_pages:
                    raise ValueError(f"Cannot split at page {split_page}. PDF only has {total_pages} pages. Please choose a page between 1 and {total_pages - 1}.")
                
                output_paths = []
                
                # First part: pages 1 to split_page
                if split_page > 0:
                    writer1 = PdfWriter()
                    
                    for i in range(split_page):
                        original_page = reader.pages[i]
                        writer1.add_page(original_page)
                    
                    # Copy metadata if available
                    if hasattr(reader, 'metadata') and reader.metadata:
                        writer1.add_metadata(reader.metadata)
                    
                    output_path1 = os.path.join(
                        self.temp_dir, 
                        f"part_1_pages_1-{split_page}_{uuid.uuid4().hex}.pdf"
                    )
                    
                    with open(output_path1, 'wb') as output_file:
                        writer1.write(output_file)
                    
                    output_paths.append(output_path1)
                    logger.info(f"Created part 1 (pages 1-{split_page}) with preserved formatting")
                
                # Second part: pages split_page+1 to end
                if split_page < total_pages:
                    writer2 = PdfWriter()
                    
                    for i in range(split_page, total_pages):
                        original_page = reader.pages[i]
                        writer2.add_page(original_page)
                    
                    # Copy metadata if available
                    if hasattr(reader, 'metadata') and reader.metadata:
                        writer2.add_metadata(reader.metadata)
                    
                    output_path2 = os.path.join(
                        self.temp_dir, 
                        f"part_2_pages_{split_page+1}-{total_pages}_{uuid.uuid4().hex}.pdf"
                    )
                    
                    with open(output_path2, 'wb') as output_file:
                        writer2.write(output_file)
                    
                    output_paths.append(output_path2)
                    logger.info(f"Created part 2 (pages {split_page+1}-{total_pages}) with preserved formatting")
                
                logger.info(f"Successfully split PDF at page {split_page} into {len(output_paths)} parts with preserved formatting")
                return output_paths
                
        except Exception as e:
            logger.error(f"Error splitting PDF at page {split_page}: {str(e)}")
            raise
    
    async def split_pdf(self, pdf_path: str, pages: Optional[str] = None) -> List[str]:
        """Split PDF into pages or extract specific pages with preserved formatting"""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                total_pages = len(reader.pages)
                
                if pages:
                    # Parse page ranges (e.g., "1-3,5,7-9")
                    try:
                        page_numbers = self._parse_page_ranges(pages, total_pages)
                    except Exception as e:
                        logger.error(f"Error parsing page ranges '{pages}': {str(e)}")
                        # Fallback to splitting all pages
                        page_numbers = list(range(1, total_pages + 1))
                else:
                    # Split all pages
                    page_numbers = list(range(1, total_pages + 1))
                
                output_paths = []
                
                if len(page_numbers) == total_pages and not pages:
                    # Split into individual pages with preserved formatting
                    for i, page_num in enumerate(page_numbers):
                        writer = PdfWriter()
                        
                        # Copy the page with all its properties
                        original_page = reader.pages[page_num - 1]
                        
                        # Preserve page properties and annotations
                        writer.add_page(original_page)
                        
                        # Copy metadata if available
                        if hasattr(reader, 'metadata') and reader.metadata:
                            writer.add_metadata(reader.metadata)
                        
                        output_path = os.path.join(
                            self.temp_dir, 
                            f"page_{page_num}_{uuid.uuid4().hex}.pdf"
                        )
                        
                        with open(output_path, 'wb') as output_file:
                            writer.write(output_file)
                        
                        output_paths.append(output_path)
                        
                        logger.info(f"Created individual page {page_num} with preserved formatting")
                else:
                    # Extract specific pages into one file with preserved formatting
                    writer = PdfWriter()
                    
                    for page_num in page_numbers:
                        # Validate page number
                        if page_num < 1 or page_num > total_pages:
                            logger.warning(f"Skipping invalid page number {page_num} (PDF has {total_pages} pages)")
                            continue
                            
                        original_page = reader.pages[page_num - 1]
                        writer.add_page(original_page)
                    
                    # Copy metadata if available
                    if hasattr(reader, 'metadata') and reader.metadata:
                        writer.add_metadata(reader.metadata)
                    
                    output_path = os.path.join(
                        self.temp_dir, 
                        f"extracted_pages_{uuid.uuid4().hex}.pdf"
                    )
                    
                    with open(output_path, 'wb') as output_file:
                        writer.write(output_file)
                    
                    output_paths.append(output_path)
                    
                    logger.info(f"Created extracted pages file with preserved formatting")
                
                if not output_paths:
                    raise ValueError("No valid pages were extracted. Please check your page range selection.")
                
                logger.info(f"Successfully split PDF into {len(output_paths)} files with preserved formatting")
                return output_paths
                
        except Exception as e:
            logger.error(f"Error splitting PDF: {str(e)}")
            raise
    
    async def compress_pdf(self, pdf_path: str) -> str:
        """Advanced PDF compression with multiple optimization techniques"""
        try:
            # Get original file size for comparison
            original_size = os.path.getsize(pdf_path)
            logger.info(f"Starting compression of PDF ({original_size} bytes)")
            
            # Try PyMuPDF compression first (most effective)
            output_path = os.path.join(self.temp_dir, f"compressed_{uuid.uuid4().hex}.pdf")
            
            try:
                # Method 1: PyMuPDF advanced compression
                compressed_path = await self._compress_with_pymupdf(pdf_path, output_path)
                
                compressed_size = os.path.getsize(compressed_path)
                ratio = (1 - compressed_size / original_size) * 100
                
                # If compression achieved good results (>2% reduction), use it
                if ratio > 2:
                    logger.info(f"PyMuPDF compression successful: {ratio:.1f}% reduction (from {original_size} to {compressed_size} bytes)")
                    return compressed_path
                else:
                    logger.info(f"PyMuPDF compression minimal ({ratio:.1f}%), trying enhanced PyPDF2 method")
                    os.remove(compressed_path)  # Clean up
                    
            except Exception as e:
                logger.warning(f"PyMuPDF compression failed: {str(e)}, falling back to PyPDF2")
            
            try:
                # Method 2: Enhanced PyPDF2 compression with optimization
                compressed_path = await self._compress_with_pypdf2_enhanced(pdf_path, output_path)
                
                compressed_size = os.path.getsize(compressed_path)
                ratio = (1 - compressed_size / original_size) * 100
                
                logger.info(f"Enhanced PyPDF2 compression completed: {ratio:.1f}% reduction (from {original_size} to {compressed_size} bytes)")
                return compressed_path
                
            except Exception as e:
                logger.warning(f"Enhanced PyPDF2 compression failed: {str(e)}, using basic compression")
            
            # Method 3: Basic PyPDF2 compression (fallback)
            compressed_path = await self._compress_with_basic_pypdf2(pdf_path, output_path)
            
            compressed_size = os.path.getsize(compressed_path)
            ratio = (1 - compressed_size / original_size) * 100
            
            logger.info(f"Basic PyPDF2 compression completed: {ratio:.1f}% reduction (from {original_size} to {compressed_size} bytes)")
            return compressed_path
            
        except Exception as e:
            logger.error(f"Error compressing PDF: {str(e)}")
            raise
    
    async def _compress_with_basic_pypdf2(self, input_path: str, output_path: str) -> str:
        """Basic PyPDF2 compression - guaranteed to work"""
        try:
            with open(input_path, 'rb') as file:
                reader = PdfReader(file)
                writer = PdfWriter()
                
                # Copy metadata
                if hasattr(reader, 'metadata') and reader.metadata:
                    writer.add_metadata(reader.metadata)
                
                # Process each page with basic compression
                for page in reader.pages:
                    # Apply content stream compression
                    page.compress_content_streams()
                    writer.add_page(page)
                
                # Write the compressed PDF
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                
                return output_path
                
        except Exception as e:
            logger.error(f"Basic PyPDF2 compression failed: {str(e)}")
            raise
    
    async def _compress_with_pymupdf(self, input_path: str, output_path: str) -> str:
        """Compress PDF using PyMuPDF with advanced optimization"""
        try:
            # Open the PDF
            doc = fitz.open(input_path)
            
            # Apply compression settings compatible with PyMuPDF version
            compression_options = {
                "garbage": 4,        # Remove unused objects (0-4, 4 is most aggressive)
                "clean": True,       # Clean up document structure
                "deflate": True,     # Use deflate compression
                "deflate_images": True,  # Compress images
                "deflate_fonts": True,   # Compress fonts
                "ascii": False,      # Don't force ASCII encoding (saves space)
                "linear": False,     # Don't linearize (saves space)
                "pretty": False,     # Don't pretty-print (saves space)
                "encryption": fitz.PDF_ENCRYPT_NONE,  # No encryption
                "permissions": -1,   # All permissions
            }
            
            # Additional image compression for each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Get all images on the page
                image_list = page.get_images(full=True)
                
                for img_index, img in enumerate(image_list):
                    try:
                        # Extract image
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # Only compress if it's a reasonably large image
                        if len(image_bytes) > 10000:  # 10KB threshold
                            # Compress image using Pillow
                            from PIL import Image
                            import io
                            
                            # Open image
                            image = Image.open(io.BytesIO(image_bytes))
                            
                            # Resize if too large (max 1920x1920 for document quality)
                            max_size = 1920
                            if image.width > max_size or image.height > max_size:
                                try:
                                    image.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
                                except AttributeError:
                                    # Fallback for older Pillow versions
                                    image.thumbnail((max_size, max_size), Image.LANCZOS)
                            
                            # Compress and save
                            output_buffer = io.BytesIO()
                            if image_ext.lower() in ['jpg', 'jpeg']:
                                image.save(output_buffer, format='JPEG', quality=85, optimize=True)
                            elif image_ext.lower() == 'png':
                                # Convert PNG to JPEG if it doesn't need transparency
                                if image.mode in ('RGBA', 'LA'):
                                    # Keep as PNG but optimize
                                    image.save(output_buffer, format='PNG', optimize=True)
                                else:
                                    # Convert to JPEG for better compression
                                    if image.mode != 'RGB':
                                        image = image.convert('RGB')
                                    image.save(output_buffer, format='JPEG', quality=85, optimize=True)
                            else:
                                # For other formats, try to convert to JPEG
                                if image.mode != 'RGB':
                                    image = image.convert('RGB')
                                image.save(output_buffer, format='JPEG', quality=85, optimize=True)
                            
                            # Replace image in PDF if compression was effective
                            compressed_bytes = output_buffer.getvalue()
                            if len(compressed_bytes) < len(image_bytes) * 0.9:  # At least 10% reduction
                                doc.update_stream(xref, compressed_bytes)
                                logger.debug(f"Compressed image {img_index} on page {page_num}: {len(image_bytes)} -> {len(compressed_bytes)} bytes")
                            
                    except Exception as img_error:
                        logger.debug(f"Could not compress image {img_index} on page {page_num}: {str(img_error)}")
                        continue
            
            # Save with compression options (removed incompatible options)
            doc.save(output_path, **compression_options)
            doc.close()
            
            return output_path
            
        except Exception as e:
            logger.error(f"PyMuPDF compression failed: {str(e)}")
            raise
    
    async def _compress_with_pypdf2_enhanced(self, input_path: str, output_path: str) -> str:
        """Enhanced PyPDF2 compression with optimization"""
        try:
            with open(input_path, 'rb') as file:
                reader = PdfReader(file)
                writer = PdfWriter()
                
                # Copy metadata
                if hasattr(reader, 'metadata') and reader.metadata:
                    writer.add_metadata(reader.metadata)
                
                # Process each page with compression
                for page in reader.pages:
                    # Apply content stream compression
                    page.compress_content_streams()
                    
                    # Add the page to writer
                    writer.add_page(page)
                
                # Note: compress_identical_objects() is not available in all PyPDF2 versions
                # The compression is already applied through page.compress_content_streams()
                
                # Write the compressed PDF
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                
                return output_path
                
        except Exception as e:
            logger.error(f"Enhanced PyPDF2 compression failed: {str(e)}")
            raise
    
    async def pdf_to_word(self, pdf_path: str) -> str:
        """Convert PDF to Word document with multiple methods"""
        try:
            # Get original file size for logging
            original_size = os.path.getsize(pdf_path)
            logger.info(f"Starting PDF to Word conversion ({original_size} bytes)")
            
            # Generate output path
            output_path = os.path.join(self.temp_dir, f"converted_{uuid.uuid4().hex}.docx")
            
            # Method 1: Try pdf2docx (most accurate for complex documents)
            try:
                cv = Converter(pdf_path)
                # Convert without specifying end parameter to avoid None issues
                cv.convert(output_path, start=0)
                cv.close()
                
                # Verify the output file was created and has content
                if os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
                    logger.info(f"Successfully converted PDF to Word using pdf2docx: {output_path}")
                    return output_path
                else:
                    logger.warning("pdf2docx conversion produced empty or invalid file, trying fallback")
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    
            except Exception as e:
                logger.warning(f"pdf2docx conversion failed: {str(e)}, trying fallback methods")
                if os.path.exists(output_path):
                    os.remove(output_path)
            
            # Method 2: Try PyMuPDF + python-docx (better text extraction)
            try:
                return await self._pdf_to_word_with_pymupdf(pdf_path, output_path)
            except Exception as e:
                logger.warning(f"PyMuPDF conversion failed: {str(e)}, using basic fallback")
            
            # Method 3: Basic fallback with PyPDF2
            return await self._pdf_to_word_fallback(pdf_path)
            
        except Exception as e:
            logger.error(f"Error converting PDF to Word: {str(e)}")
            raise
    
    async def _pdf_to_word_with_pymupdf(self, pdf_path: str, output_path: str) -> str:
        """Convert PDF to Word using PyMuPDF for better text extraction"""
        try:
            import fitz  # PyMuPDF
            from docx import Document
            from docx.shared import Inches
            
            # Open PDF with PyMuPDF
            doc_pdf = fitz.open(pdf_path)
            
            # Create Word document
            doc_word = Document()
            doc_word.add_heading('Converted from PDF', 0)
            
            # Process each page
            for page_num in range(len(doc_pdf)):
                page = doc_pdf[page_num]
                
                # Add page heading
                if len(doc_pdf) > 1:
                    doc_word.add_heading(f'Page {page_num + 1}', level=1)
                
                # Extract text blocks (better formatting preservation)
                blocks = page.get_text("dict")
                
                for block in blocks.get("blocks", []):
                    if "lines" in block:
                        # Text block
                        paragraph_text = ""
                        for line in block["lines"]:
                            for span in line.get("spans", []):
                                text = span.get("text", "").strip()
                                if text:
                                    paragraph_text += text + " "
                        
                        if paragraph_text.strip():
                            # Detect if it might be a heading (larger font, bold, etc.)
                            first_span = block["lines"][0]["spans"][0] if block["lines"] and block["lines"][0].get("spans") else {}
                            font_size = first_span.get("size", 12)
                            font_flags = first_span.get("flags", 0)
                            
                            if font_size > 14 or (font_flags & 2**4):  # Large font or bold
                                doc_word.add_heading(paragraph_text.strip(), level=2)
                            else:
                                doc_word.add_paragraph(paragraph_text.strip())
                
                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc_pdf.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # Save image temporarily
                        temp_img_path = os.path.join(self.temp_dir, f"temp_img_{page_num}_{img_index}.{image_ext}")
                        with open(temp_img_path, "wb") as img_file:
                            img_file.write(image_bytes)
                        
                        # Add image to Word document
                        try:
                            doc_word.add_picture(temp_img_path, width=Inches(4.0))
                            logger.debug(f"Added image to Word document: page {page_num}, image {img_index}")
                        except Exception as img_error:
                            logger.debug(f"Could not add image to Word: {str(img_error)}")
                        finally:
                            # Clean up temp image
                            if os.path.exists(temp_img_path):
                                os.remove(temp_img_path)
                                
                    except Exception as img_error:
                        logger.debug(f"Could not extract image {img_index} from page {page_num}: {str(img_error)}")
                        continue
                
                # Add page break (except for last page)
                if page_num < len(doc_pdf) - 1:
                    doc_word.add_page_break()
            
            doc_pdf.close()
            
            # Save Word document
            doc_word.save(output_path)
            
            logger.info(f"Successfully converted PDF to Word using PyMuPDF: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"PyMuPDF PDF to Word conversion failed: {str(e)}")
            raise
    
    async def word_to_pdf(self, word_path: str) -> str:
        """Convert Word document to PDF"""
        try:
            # Read Word document
            doc = Document(word_path)
            
            # Generate output path
            output_path = os.path.join(self.temp_dir, f"converted_{uuid.uuid4().hex}.pdf")
            
            # Create PDF using reportlab
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            
            # Create PDF
            doc_pdf = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Extract text from Word document and add to PDF
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    p = Paragraph(paragraph.text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            doc_pdf.build(story)
            
            logger.info(f"Successfully converted Word to PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error converting Word to PDF: {str(e)}")
            raise
    
    def _parse_page_ranges(self, pages: str, total_pages: int) -> List[int]:
        """Parse page range string like '1-3,5,7-9' into list of page numbers"""
        page_numbers = []
        
        for part in pages.split(','):
            part = part.strip()
            if '-' in part:
                # Range like "1-3"
                start, end = map(int, part.split('-'))
                page_numbers.extend(range(start, end + 1))
            else:
                # Single page like "5"
                page_numbers.append(int(part))
        
        # Filter valid page numbers
        valid_pages = [p for p in page_numbers if 1 <= p <= total_pages]
        return sorted(list(set(valid_pages)))  # Remove duplicates and sort
    
    async def _pdf_to_word_fallback(self, pdf_path: str) -> str:
        """Enhanced fallback method for PDF to Word conversion using PyPDF2"""
        try:
            # Extract text from PDF
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                
                # Create Word document
                doc = Document()
                doc.add_heading('Converted from PDF', 0)
                
                # Add document info
                if hasattr(reader, 'metadata') and reader.metadata:
                    metadata = reader.metadata
                    info_paragraph = doc.add_paragraph()
                    info_paragraph.add_run('Document Information:').bold = True
                    if metadata.get('/Title'):
                        doc.add_paragraph(f"Title: {metadata['/Title']}")
                    if metadata.get('/Author'):
                        doc.add_paragraph(f"Author: {metadata['/Author']}")
                    if metadata.get('/Subject'):
                        doc.add_paragraph(f"Subject: {metadata['/Subject']}")
                    doc.add_paragraph()  # Add space
                
                # Process each page
                for i, page in enumerate(reader.pages):
                    try:
                        # Extract text
                        text = page.extract_text()
                        
                        if text and text.strip():
                            # Add page heading for multi-page documents
                            if len(reader.pages) > 1:
                                doc.add_heading(f'Page {i + 1}', level=1)
                            
                            # Split text into paragraphs (by double newlines or long single newlines)
                            paragraphs = []
                            current_paragraph = ""
                            
                            lines = text.split('\n')
                            for line in lines:
                                line = line.strip()
                                if not line:  # Empty line
                                    if current_paragraph.strip():
                                        paragraphs.append(current_paragraph.strip())
                                        current_paragraph = ""
                                else:
                                    if current_paragraph:
                                        current_paragraph += " " + line
                                    else:
                                        current_paragraph = line
                            
                            # Add remaining paragraph
                            if current_paragraph.strip():
                                paragraphs.append(current_paragraph.strip())
                            
                            # Add paragraphs to document
                            for paragraph_text in paragraphs:
                                if len(paragraph_text) > 200 or not paragraph_text.endswith('.'):
                                    # Long paragraph or incomplete sentence - likely body text
                                    doc.add_paragraph(paragraph_text)
                                elif len(paragraph_text) < 100 and paragraph_text.isupper():
                                    # Short, all caps text - likely a heading
                                    doc.add_heading(paragraph_text, level=2)
                                elif len(paragraph_text) < 80 and not paragraph_text.endswith('.'):
                                    # Short text without period - likely a heading or title
                                    doc.add_heading(paragraph_text, level=2)
                                else:
                                    # Regular paragraph
                                    doc.add_paragraph(paragraph_text)
                        else:
                            # No text extracted from this page
                            if len(reader.pages) > 1:
                                doc.add_heading(f'Page {i + 1}', level=1)
                            doc.add_paragraph('[No text content could be extracted from this page]')
                            
                    except Exception as page_error:
                        logger.warning(f"Error extracting text from page {i + 1}: {str(page_error)}")
                        if len(reader.pages) > 1:
                            doc.add_heading(f'Page {i + 1}', level=1)
                        doc.add_paragraph(f'[Error extracting content from page {i + 1}: {str(page_error)}]')
                
                # Generate output path
                output_path = os.path.join(self.temp_dir, f"converted_fallback_{uuid.uuid4().hex}.docx")
                doc.save(output_path)
                
                logger.info(f"Successfully converted PDF to Word using fallback method: {output_path}")
                return output_path
                
        except Exception as e:
            logger.error(f"Error in fallback PDF to Word conversion: {str(e)}")
            raise